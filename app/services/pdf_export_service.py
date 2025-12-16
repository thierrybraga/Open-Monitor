"""
Serviço para exportação de relatórios em PDF
Suporta múltiplas bibliotecas: WeasyPrint, ReportLab, e fallback para HTML
"""

import os
import io
import base64
import logging
from datetime import datetime
from typing import Optional, Dict, Any, List, Union
from pathlib import Path

from flask import current_app, render_template, url_for
from jinja2 import Template

logger = logging.getLogger(__name__)

class PDFExportService:
    """Serviço para exportação de relatórios em PDF"""
    
    def __init__(self):
        self.pdf_engine = self._detect_pdf_engine()
        self.temp_dir = None

    def export_to_pdf(self, report, template_name: Optional[str] = None, options: Optional[Dict[str, Any]] = None) -> str:
        """Wrapper compatível com controller: retorna somente o caminho do arquivo."""
        result = self.export_report(report, format_type='pdf', template_name=template_name, options=options)
        return result.get('filepath')

    def export_to_docx(self, report, template_name: Optional[str] = None, options: Optional[Dict[str, Any]] = None) -> str:
        """Wrapper compatível com controller: retorna somente o caminho do arquivo."""
        result = self.export_report(report, format_type='docx', template_name=template_name, options=options)
        return result.get('filepath')
        
    def _detect_pdf_engine(self) -> str:
        """Detecta qual engine PDF está disponível, preferindo configuração/ambiente."""
        import os
        preferred = os.getenv('PDF_ENGINE', 'reportlab').lower()

        candidates = []
        if preferred in ('weasyprint', 'reportlab', 'pdfkit'):
            candidates.append(preferred)
        # completar ordem de tentativa
        for c in ('weasyprint', 'reportlab', 'pdfkit'):
            if c not in candidates:
                candidates.append(c)

        # tentar importar na ordem
        for c in candidates:
            try:
                if c == 'weasyprint':
                    import weasyprint  # noqa: F401
                elif c == 'reportlab':
                    import reportlab  # noqa: F401
                elif c == 'pdfkit':
                    import pdfkit  # noqa: F401
                logger.info(f"Engine PDF selecionada: {c}")
                return c
            except Exception as e:
                logger.warning(f"Engine {c} indisponível, tentando próxima. Motivo: {e}")

        logger.warning("Nenhum engine PDF encontrado. Usando fallback HTML.")
        return 'html'
    
    def _ensure_temp_dir(self):
        """Inicializa o diretório temporário se necessário"""
        if self.temp_dir is None:
            from flask import current_app
            # Garantir caminho absoluto sob a raiz do projeto (um nível acima de app.root_path)
            base_dir = Path(current_app.root_path).parent
            temp_cfg = current_app.config.get('TEMP_DIR', 'temp')
            self.temp_dir = (base_dir / temp_cfg).resolve()
            self.temp_dir.mkdir(exist_ok=True)
    
    def export_report(self, report, format_type: str = 'pdf', 
                     template_name: Optional[str] = None,
                     options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Exporta um relatório para PDF ou outros formatos
        
        Args:
            report: Objeto do relatório
            format_type: Tipo de formato ('pdf', 'html', 'docx')
            template_name: Nome do template específico
            options: Opções de exportação
            
        Returns:
            Dict com dados do arquivo exportado
        """
        try:
            options = options or {}
            
            # Preparar dados do relatório
            report_data = self._prepare_report_data(report)
            
            # Determinar template
            if not template_name:
                template_name = self._get_template_for_report_type(report.report_type)
            
            if format_type.lower() == 'pdf':
                return self._export_to_pdf(report_data, template_name, options)
            elif format_type.lower() == 'html':
                return self._export_to_html(report_data, template_name, options)
            elif format_type.lower() == 'docx':
                return self._export_to_docx(report_data, template_name, options)
            else:
                raise ValueError(f"Formato não suportado: {format_type}")
                
        except Exception as e:
            logger.error(f"Erro ao exportar relatório {report.id}: {str(e)}")
            raise
    
    def _prepare_report_data(self, report) -> Dict[str, Any]:
        """Prepara os dados do relatório para exportação"""
        return {
            'report': report,
            'export_date': datetime.now(),
            'charts_data': self._prepare_charts_data(report),
            'summary_stats': self._calculate_summary_stats(report),
            'formatted_content': self._format_content_for_pdf(report.content),
            'company_info': self._get_company_info(),
            'export_metadata': {
                'generated_by': 'Open Monitor',
                'version': current_app.config.get('VERSION', '1.0.0'),
                'export_engine': self.pdf_engine
            }
        }
    
    def _prepare_charts_data(self, report) -> Dict[str, Any]:
        """Prepara dados dos gráficos para exportação"""
        content = report.content or {}
        if not isinstance(content, dict) or 'charts' not in content:
            return {}

        charts_data = {}

        try:
            # Gráfico de distribuição CVSS
            if 'vulnerabilities' in content:
                charts_data['cvss_distribution'] = self._generate_cvss_chart_data(report)

            # Gráfico de tendências
            if 'trends' in content or 'charts' in content:
                charts_data['trends'] = self._generate_trends_chart_data(report)

            # Gráfico de ativos por risco
            if 'assets' in content:
                charts_data['assets_risk'] = self._generate_assets_risk_chart_data(report)

        except Exception as e:
            logger.warning(f"Erro ao preparar dados dos gráficos: {str(e)}")

        return charts_data
    
    def _calculate_summary_stats(self, report) -> Dict[str, Any]:
        """Calcula estatísticas resumidas do relatório"""
        stats = {
            'total_vulnerabilities': 0,
            'critical_count': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'affected_assets': 0,
            'avg_cvss': 0.0,
            'risk_score': 0.0
        }
        
        try:
            content = report.content or {}
            if isinstance(content, dict) and 'vulnerabilities' in content:
                vulns = content.get('vulnerabilities', {})
                if isinstance(vulns, dict):
                    vuln_list = vulns.get('details', [])
                    stats['total_vulnerabilities'] = len(vuln_list)
                    
                    cvss_scores = []
                    for vuln in vuln_list:
                        severity = vuln.get('severity', '').lower()
                        if severity == 'critical':
                            stats['critical_count'] += 1
                        elif severity == 'high':
                            stats['high_count'] += 1
                        elif severity == 'medium':
                            stats['medium_count'] += 1
                        elif severity == 'low':
                            stats['low_count'] += 1
                            
                        if vuln.get('cvss_score'):
                            cvss_scores.append(float(vuln['cvss_score']))
                    
                    if cvss_scores:
                        stats['avg_cvss'] = round(sum(cvss_scores) / len(cvss_scores), 1)
                    
                    # Calcular score de risco
                    stats['risk_score'] = self._calculate_risk_score(stats)
                    
        except Exception as e:
            logger.warning(f"Erro ao calcular estatísticas: {str(e)}")
            
        return stats
    
    def _calculate_risk_score(self, stats: Dict[str, Any]) -> float:
        """Calcula um score de risco baseado nas vulnerabilidades"""
        try:
            # Pesos por severidade
            weights = {'critical': 10, 'high': 7, 'medium': 4, 'low': 1}
            
            total_score = (
                stats['critical_count'] * weights['critical'] +
                stats['high_count'] * weights['high'] +
                stats['medium_count'] * weights['medium'] +
                stats['low_count'] * weights['low']
            )
            
            # Normalizar para escala 0-100
            max_possible = stats['total_vulnerabilities'] * weights['critical']
            if max_possible > 0:
                return round((total_score / max_possible) * 100, 1)
            
        except Exception as e:
            logger.warning(f"Erro ao calcular risk score: {str(e)}")
            
        return 0.0
    
    def _format_content_for_pdf(self, content) -> Dict[str, Any]:
        """Formata o conteúdo para melhor exibição em PDF"""
        if not content or not isinstance(content, dict):
            return {}

        formatted: Dict[str, Any] = {}

        try:
            # Formatar vulnerabilidades
            if 'vulnerabilities' in content:
                formatted['vulnerabilities'] = self._format_vulnerabilities_for_pdf(content.get('vulnerabilities'))

            # Formatar análise de impacto (fallback para BIA)
            if 'business_impact' in content and isinstance(content.get('business_impact'), dict):
                formatted['business_impact'] = self._format_business_impact_for_pdf(content.get('business_impact'))
            elif 'bia_analysis' in content:
                bia = content.get('bia_analysis') or {}
                # `bia_analysis` no template de visualização é markdown em `content`.
                # Para o PDF padrão, usamos como impacto operacional, mantendo outras áreas como "Não avaliado".
                bia_md = bia.get('content') if isinstance(bia, dict) else None
                formatted['business_impact'] = self._format_business_impact_for_pdf({
                    'operational_impact': bia_md or 'Não avaliado'
                })

            # Formatar recomendações (fallback para plano de remediação)
            if 'recommendations' in content and isinstance(content.get('recommendations'), list):
                formatted['recommendations'] = self._format_recommendations_for_pdf(content.get('recommendations'))
            elif 'remediation_plan' in content:
                rem = content.get('remediation_plan') or {}
                rem_md = rem.get('content') if isinstance(rem, dict) else None
                # Criar uma recomendação única baseada no plano de remediação
                formatted['recommendations'] = self._format_recommendations_for_pdf([
                    {
                        'title': 'Plano de Remediação',
                        'description': rem_md or 'Sem detalhes',
                        'priority': 'medium',
                        'effort': 'Não estimado',
                        'timeline': 'Não definido'
                    }
                ])

        except Exception as e:
            logger.warning(f"Erro ao formatar conteúdo: {str(e)}")

        return formatted
    
    def _get_company_info(self) -> Dict[str, Any]:
        """Obtém informações da empresa para o cabeçalho do PDF"""
        return {
            'name': current_app.config.get('COMPANY_NAME', 'Open Monitor'),
            'logo_url': current_app.config.get('COMPANY_LOGO', ''),
            'address': current_app.config.get('COMPANY_ADDRESS', ''),
            'website': current_app.config.get('COMPANY_WEBSITE', ''),
            'contact': current_app.config.get('COMPANY_CONTACT', '')
        }
    
    def _get_template_for_report_type(self, report_type: str) -> str:
        """Determina o template baseado no tipo de relatório"""
        template_mapping = {
            'executive': 'reports/pdf/executive_pdf.html',
            'technical': 'reports/pdf/technical_pdf.html',
            'estudo_tecnico': 'reports/pdf/technical_study_pdf.html',
            'compliance': 'reports/pdf/compliance_pdf.html',
            'pentest': 'reports/pdf/pentest_pdf.html',
            'bia': 'reports/pdf/bia_pdf.html',
            'kpi': 'reports/pdf/kpi_pdf.html'
        }
        
        return template_mapping.get(report_type.value if hasattr(report_type, 'value') else report_type, 
                                  'reports/pdf/default_pdf.html')
    
    def _export_to_pdf(self, report_data: Dict[str, Any], 
                      template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta para PDF usando engine com fallback em cascata."""
        from flask import current_app
        import os

        # Ordem preferida: configuração -> engines restantes
        preferred = None
        try:
            preferred = current_app.config.get('PDF_ENGINE')
        except Exception:
            preferred = None
        if not preferred:
            preferred = os.getenv('PDF_ENGINE', self.pdf_engine)

        engines = []
        if preferred in ('weasyprint', 'reportlab', 'pdfkit'):
            engines.append(preferred)
        for e in ('weasyprint', 'reportlab', 'pdfkit'):
            if e not in engines:
                engines.append(e)

        last_error = None
        for engine in engines:
            try:
                logger.info(f"Tentando exportar PDF com engine: {engine}")
                if engine == 'weasyprint':
                    return self._export_with_weasyprint(report_data, template_name, options)
                elif engine == 'reportlab':
                    return self._export_with_reportlab(report_data, template_name, options)
                elif engine == 'pdfkit':
                    return self._export_with_pdfkit(report_data, template_name, options)
            except Exception as e:
                logger.error(f"Falha ao exportar com {engine}: {e}")
                last_error = e
                continue

        logger.warning("Todas as engines de PDF falharam. Usando fallback HTML.")
        return self._export_to_html(report_data, template_name, options)
    
    def _export_with_weasyprint(self, report_data: Dict[str, Any], 
                               template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta usando WeasyPrint"""
        try:
            import weasyprint
            from weasyprint import HTML, CSS
            
            # Renderizar HTML
            html_content = render_template(template_name, **report_data)
            
            # CSS customizado para PDF
            css_content = self._get_pdf_css()
            
            # Configurações do WeasyPrint
            base_url = current_app.config.get('BASE_URL', 'http://localhost:4443')
            
            # Gerar PDF
            html_doc = HTML(string=html_content, base_url=base_url)
            css_doc = CSS(string=css_content)
            
            pdf_buffer = io.BytesIO()
            html_doc.write_pdf(pdf_buffer, stylesheets=[css_doc])
            pdf_content = pdf_buffer.getvalue()
            
            # Salvar arquivo temporário
            self._ensure_temp_dir()
            filename = f"report_{report_data['report'].id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.temp_dir / filename
            
            with open(filepath, 'wb') as f:
                f.write(pdf_content)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'content': pdf_content,
                'size': len(pdf_content),
                'mime_type': 'application/pdf'
            }
            
        except Exception as e:
            logger.error(f"Erro no WeasyPrint: {str(e)}")
            raise
    
    def _export_with_reportlab(self, report_data: Dict[str, Any], 
                              template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta usando ReportLab"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Buffer para o PDF
            pdf_buffer = io.BytesIO()
            
            # Configurar documento
            page_size = A4 if options.get('page_size') == 'A4' else letter
            doc = SimpleDocTemplate(pdf_buffer, pagesize=page_size,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#2c3e50')
            )
            
            # Construir conteúdo
            story = []
            report = report_data['report']
            
            # Título
            story.append(Paragraph(report.title, title_style))
            story.append(Spacer(1, 12))
            
            # Informações básicas
            basic_info = [
                ['Tipo:', report.report_type.value.replace('_', ' ').title()],
                ['Data de Criação:', report.created_at.strftime('%d/%m/%Y %H:%M')],
                ['Escopo:', report.scope.value.replace('_', ' ').title()],
                ['Status:', report.status.value.title()]
            ]
            
            basic_table = Table(basic_info, colWidths=[2*inch, 4*inch])
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(basic_table)
            story.append(Spacer(1, 20))
            
            # Estatísticas
            stats = report_data['summary_stats']
            if stats['total_vulnerabilities'] > 0:
                story.append(Paragraph("Resumo de Vulnerabilidades", styles['Heading2']))
                
                vuln_data = [
                    ['Severidade', 'Quantidade'],
                    ['Críticas', str(stats['critical_count'])],
                    ['Altas', str(stats['high_count'])],
                    ['Médias', str(stats['medium_count'])],
                    ['Baixas', str(stats['low_count'])],
                    ['Total', str(stats['total_vulnerabilities'])]
                ]
                
                vuln_table = Table(vuln_data, colWidths=[3*inch, 2*inch])
                vuln_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(vuln_table)
            
            # Gerar PDF
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            
            # Salvar arquivo
            self._ensure_temp_dir()
            filename = f"report_{report.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.temp_dir / filename
            
            with open(filepath, 'wb') as f:
                f.write(pdf_content)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'content': pdf_content,
                'size': len(pdf_content),
                'mime_type': 'application/pdf'
            }
            
        except Exception as e:
            logger.error(f"Erro no ReportLab: {str(e)}")
            raise
    
    def _export_with_pdfkit(self, report_data: Dict[str, Any], 
                           template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta usando pdfkit (wkhtmltopdf)"""
        try:
            import pdfkit
            from flask import current_app
            import os
            
            # Renderizar HTML
            html_content = render_template(template_name, **report_data)
            
            # Opções do wkhtmltopdf
            pdf_options = {
                'page-size': options.get('page_size', 'A4'),
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            # Configuração do executável wkhtmltopdf, se fornecida
            wkhtml_path = None
            try:
                wkhtml_path = current_app.config.get('WKHTMLTOPDF_PATH')
            except Exception:
                wkhtml_path = None
            if not wkhtml_path:
                wkhtml_path = os.getenv('WKHTMLTOPDF_PATH')
            configuration = None
            if wkhtml_path:
                try:
                    configuration = pdfkit.configuration(wkhtmltopdf=wkhtml_path)
                    logger.info(f"Usando wkhtmltopdf em: {wkhtml_path}")
                except Exception as e:
                    logger.warning(f"Configuração wkhtmltopdf inválida ({wkhtml_path}): {e}")
            
            # Gerar PDF
            pdf_content = pdfkit.from_string(
                html_content,
                False,
                options=pdf_options,
                configuration=configuration if configuration else None
            )
            
            # Salvar arquivo
            self._ensure_temp_dir()
            filename = f"report_{report_data['report'].id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.temp_dir / filename
            
            with open(filepath, 'wb') as f:
                f.write(pdf_content)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'content': pdf_content,
                'size': len(pdf_content),
                'mime_type': 'application/pdf'
            }
            
        except Exception as e:
            logger.error(f"Erro no pdfkit: {str(e)}")
            raise
    
    def _export_to_html(self, report_data: Dict[str, Any], 
                       template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta para HTML como fallback"""
        try:
            # Renderizar HTML
            html_content = render_template(template_name, **report_data)
            
            # Salvar arquivo
            self._ensure_temp_dir()
            filename = f"report_{report_data['report'].id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            filepath = self.temp_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'content': html_content.encode('utf-8'),
                'size': len(html_content.encode('utf-8')),
                'mime_type': 'text/html'
            }
            
        except Exception as e:
            logger.error(f"Erro na exportação HTML: {str(e)}")
            raise
    
    def _export_to_docx(self, report_data: Dict[str, Any], 
                       template_name: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Exporta para DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Criar documento
            doc = Document()
            report = report_data['report']
            
            # Título
            title = doc.add_heading(report.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informações básicas
            doc.add_heading('Informações do Relatório', level=1)
            
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Tipo:', report.report_type.value.replace('_', ' ').title()),
                ('Data de Criação:', report.created_at.strftime('%d/%m/%Y %H:%M')),
                ('Escopo:', report.scope.value.replace('_', ' ').title()),
                ('Status:', report.status.value.title())
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
            
            # Estatísticas
            stats = report_data['summary_stats']
            if stats['total_vulnerabilities'] > 0:
                doc.add_heading('Resumo de Vulnerabilidades', level=1)
                
                stats_table = doc.add_table(rows=6, cols=2)
                stats_table.style = 'Table Grid'
                
                stats_data = [
                    ('Severidade', 'Quantidade'),
                    ('Críticas', str(stats['critical_count'])),
                    ('Altas', str(stats['high_count'])),
                    ('Médias', str(stats['medium_count'])),
                    ('Baixas', str(stats['low_count'])),
                    ('Total', str(stats['total_vulnerabilities']))
                ]
                
                for i, (label, value) in enumerate(stats_data):
                    stats_table.cell(i, 0).text = label
                    stats_table.cell(i, 1).text = value
            
            # Salvar arquivo
            self._ensure_temp_dir()
            filename = f"report_{report.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.temp_dir / filename
            
            doc.save(str(filepath))
            
            # Ler conteúdo para retorno
            with open(filepath, 'rb') as f:
                content = f.read()
            
            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'content': content,
                'size': len(content),
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
        except Exception as e:
            logger.error(f"Erro na exportação DOCX: {str(e)}")
            raise
    
    def _get_pdf_css(self) -> str:
        """Retorna CSS otimizado para PDF"""
        return """
        @page {
            size: A4;
            margin: 2cm;
            @top-center {
                content: "Open Monitor - Relatório de Segurança";
                font-size: 10pt;
                color: #666;
            }
            @bottom-center {
                content: "Página " counter(page) " de " counter(pages);
                font-size: 10pt;
                color: #666;
            }
        }
        
        body {
            font-family: 'Helvetica', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }
        
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            page-break-after: avoid;
        }
        
        h1 { font-size: 24pt; margin-bottom: 20pt; }
        h2 { font-size: 18pt; margin-bottom: 15pt; }
        h3 { font-size: 14pt; margin-bottom: 12pt; }
        
        .page-break {
            page-break-before: always;
        }
        
        .no-break {
            page-break-inside: avoid;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15pt;
            page-break-inside: avoid;
        }
        
        th, td {
            border: 1pt solid #ddd;
            padding: 8pt;
            text-align: left;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        
        .severity-critical { color: #dc3545; font-weight: bold; }
        .severity-high { color: #fd7e14; font-weight: bold; }
        .severity-medium { color: #ffc107; font-weight: bold; }
        .severity-low { color: #28a745; font-weight: bold; }
        
        .chart-placeholder {
            width: 100%;
            height: 300pt;
            border: 1pt solid #ddd;
            background-color: #f8f9fa;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 15pt 0;
        }
        
        .executive-summary {
            background-color: #f8f9fa;
            padding: 15pt;
            border-left: 4pt solid #007bff;
            margin-bottom: 20pt;
        }
        
        .vulnerability-item {
            border: 1pt solid #ddd;
            margin-bottom: 15pt;
            padding: 10pt;
            page-break-inside: avoid;
        }
        
        .code-block {
            background-color: #f8f9fa;
            border: 1pt solid #e9ecef;
            padding: 10pt;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
            margin: 10pt 0;
            page-break-inside: avoid;
        }
        """
    
    def _generate_cvss_chart_data(self, report) -> Dict[str, Any]:
        """Gera dados para gráfico de distribuição CVSS"""
        # Implementação simplificada - seria expandida com dados reais
        return {
            'type': 'doughnut',
            'data': {
                'labels': ['Críticas', 'Altas', 'Médias', 'Baixas'],
                'datasets': [{
                    'data': [5, 12, 8, 3],
                    'backgroundColor': ['#dc3545', '#fd7e14', '#ffc107', '#28a745']
                }]
            }
        }
    
    def _generate_trends_chart_data(self, report) -> Dict[str, Any]:
        """Gera dados para gráfico de tendências"""
        return {
            'type': 'line',
            'data': {
                'labels': ['Jan', 'Fev', 'Mar', 'Abr', 'Mai'],
                'datasets': [{
                    'label': 'Vulnerabilidades',
                    'data': [10, 15, 12, 8, 5],
                    'borderColor': '#007bff'
                }]
            }
        }
    
    def _generate_assets_risk_chart_data(self, report) -> Dict[str, Any]:
        """Gera dados para gráfico de ativos por risco"""
        return {
            'type': 'bar',
            'data': {
                'labels': ['Servidor Web', 'Database', 'Firewall', 'Workstations'],
                'datasets': [{
                    'label': 'Score de Risco',
                    'data': [85, 65, 45, 25],
                    'backgroundColor': '#dc3545'
                }]
            }
        }
    
    def _format_vulnerabilities_for_pdf(self, vulnerabilities) -> List[Dict[str, Any]]:
        """Formata vulnerabilidades para PDF"""
        if not vulnerabilities or not isinstance(vulnerabilities, dict):
            return []
            
        formatted = []
        vuln_list = vulnerabilities.get('details', [])
        
        for vuln in vuln_list:
            formatted.append({
                'title': vuln.get('title', 'Vulnerabilidade sem título'),
                'cve_id': vuln.get('cve_id', 'N/A'),
                'severity': vuln.get('severity', 'unknown'),
                'cvss_score': vuln.get('cvss_score', 0),
                'description': vuln.get('description', ''),
                'affected_assets_count': len(vuln.get('affected_assets', [])),
                'exploit_available': vuln.get('exploit_available', False)
            })
            
        return formatted
    
    def _format_business_impact_for_pdf(self, business_impact) -> Dict[str, Any]:
        """Formata análise de impacto para PDF"""
        if not business_impact:
            return {}
            
        return {
            'financial_impact': business_impact.get('financial_impact', 'Não avaliado'),
            'operational_impact': business_impact.get('operational_impact', 'Não avaliado'),
            'reputation_impact': business_impact.get('reputation_impact', 'Não avaliado'),
            'compliance_impact': business_impact.get('compliance_impact', 'Não avaliado')
        }
    
    def _format_recommendations_for_pdf(self, recommendations) -> List[Dict[str, Any]]:
        """Formata recomendações para PDF"""
        if not recommendations:
            return []
            
        formatted = []
        if isinstance(recommendations, list):
            for rec in recommendations:
                formatted.append({
                    'title': rec.get('title', 'Recomendação'),
                    'description': rec.get('description', ''),
                    'priority': rec.get('priority', 'medium'),
                    'effort': rec.get('effort', 'Não estimado'),
                    'timeline': rec.get('timeline', 'Não definido')
                })
                
        return formatted

    def cleanup_temp_files(self, older_than_hours: int = 24):
        """Remove arquivos temporários antigos"""
        try:
            self._ensure_temp_dir()
            import time
            current_time = time.time()
            
            for file_path in self.temp_dir.glob('*'):
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > (older_than_hours * 3600):
                        file_path.unlink()
                        logger.info(f"Arquivo temporário removido: {file_path}")
                        
        except Exception as e:
            logger.error(f"Erro ao limpar arquivos temporários: {str(e)}")
